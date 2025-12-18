from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import matplotlib.pyplot as plt
from datetime import date
import os
from docx2pdf import convert
from lxml.html.defs import list_tags
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

# === CONFIGURAÇÕES === #
MUNICIPIO = "Brasilândia"
DATA_RELATORIO = date.today().strftime("%d/%m/%Y")
ANO_RELATORIO = date.today().strftime("%Y")
NUM_RELATORIO = "0XX"
PROCESSO = "51.010.469-2025"

# Caminhos
MODELO = "modelo_agems_demanda.docx"
SAIDA_DOCX = f"saida/minuta_of_{MUNICIPIO.lower()}.docx"
SAIDA_PDF = f"saida/minuta_of_{MUNICIPIO.lower()}.pdf"

os.makedirs("saida", exist_ok=True)
os.makedirs("figuras", exist_ok=True)

# === ABRE O MODELO BASE === #
doc = Document(MODELO)

# === 1. INÍCIO === #
titulo = doc.add_paragraph(f"Ofício n. {NUM_RELATORIO}/DGE/AGEMS/{ANO_RELATORIO}\n")
titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = titulo.runs[0]
run.font.size = Pt(14)
run.bold = True

titulo = doc.add_paragraph(f"Campo Grande/MS, XX de XXXXXX {ANO_RELATORIO}\n")
titulo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = titulo.runs[0]
run.font.size = Pt(14)
run.bold = True

# for style in doc.styles:
# print(style.name)

# === 2. ASSUNTO === #
solicitante = f"Promotoria de Justiça da Comarca de {MUNICIPIO}, Ministério Público de Mato Grosso do Sul"
solicitante2 = "Promotoria"
nome_resp_solicitante = "Adriano Barrozo da Silva"
cargo_resp_solicitante = "Promotor de Justiça"
comarca = "Comarca de Brasilândia"
num_oficio_demandado = "0629/2025/PJ/BRS"
num_oficio_env_conce = "5103/2025/DGE"
num_oficio_resp_conc = "ENERGISAMS/DTEC-ANEEL/N°070/202"
localidade = "Assentamento Mutum"
localidade2 = "assentamento"
pronom_tratamento = "Senhor Promotor"
data_oficio = "04 de novembro de 2025"
noticia_fato = "09.2025.00008077-8"
concessionaria = "Energisa Mato Grosso do Sul – EMS"
concessionaria2 = "EMS"
concessionaria3 = "ENERGISA MS"
conjunto1 = "Água Clara"
conjunto2 = conjunto1
ano_inicio_analise = "2022"
mes_inicio_analise = "janeiro"
ano_final_analise = "2025"
mes_final_analise = "novembro"
nivel_tensao_13_8 = "13,8 kV"
nivel_tensao_34_5 = "34,5 kV"
figura = 1
tabela = 1

def destacar_palvras(paragrafo_obj, texto, palavras_destaque, highlight=True, underline=True):
    texto = paragrafo_obj.text

    paragrafo_obj.clear()  # limpa o parágrafo

    pattern = re.compile("(" + "|".join(re.escape(w) for w in palavras_destaque) + ")")
    partes = re.split(pattern, texto)

    for parte in partes:
        if not parte:
            continue
        run = paragrafo_obj.add_run(parte)
        if parte in palavras_destaque:
            if highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if underline:
                run.underline = True

assunto = (
    f"""Assunto: Ofício n° {num_oficio_demandado} - {solicitante}.
Processo: {PROCESSO}
        """
)

a = doc.add_paragraph(assunto)
a.alignment = WD_ALIGN_PARAGRAPH.LEFT
a.style = "Normal"
run = a.runs[0]
doc.add_paragraph("")

# === 4. CORPO === #

paragrafo = [
    f"{pronom_tratamento},",

    f"Reportamo-nos ao Ofício em epígrafe, acerca do desempenho da {concessionaria3} em relação ao fornecimento de energia no "
    f"{localidade} localizado na área rural de Ribas do Rio Pardo.",

    f"Sobre esse assunto, em primeiro lugar, salientamos que a Agência Nacional de Energia Elétrica - ANEEL monitora a "
    f"qualidade da distribuição de energia elétrica por meio dos seguintes indicadores de continuidade de fornecimento: DEC"
    f"(Duração Equivalente de Interrupção por Unidade Consumidora) e FEC (Frequência Equivalente de Interrupção por Unidade Consumidora).",

    f"Esses indicadores refletem o desempenho das Distribuidoras na prestação do serviço público e são avaliados para Conjuntos"
    f"Elétricos, que são agrupamentos de alimentadores derivados de uma subestação que agrupam segmentos contínuos de uma "
    f"área de concessão. Os estudos de qualidade e conformidade da rede elétrica realizados pela ANEEL são disponibilizados "
    f"para os Conjuntos Elétricos",

    f"O alimentador que atende o {localidade}, identificado como Água Clara 01, pertence ao Conjunto Elétrico Água Clara "
    f"formado pelo agrupamento de alimentadores da Subestação Água Clara. Este conjunto apresenta os indicadores de continuidade "
    f"DEC e FEC dentro dos limites regulados pela ANEEL.",

    f"Para restringir a análise às unidades consumidoras conectadas ao alimentador que atende o {localidade}, nossos fiscais"
    f"analisaram, no período {ano_inicio_analise} a {ano_final_analise}, os indicadores individuais de continidade "
    f"DIC (Duração de Interrupção Individual por Unidade Consumidora), FIC (Frequência de Interrupção Individual por Unidade "
    f"Consumidora) e DMIC (Duração Máxima de Interrupção Contínua por Unidade Consumidora) que são destacados na fatura de energia "
    f"elétrica do consumidor.",

    f"Esses indicadores representam, para o consumidor, a qualidade dos serviços prestados pela Distribuidora e mensuram "
    f"a duração e a frequência das interrupções ocorridas em sua unidade consumidora. Os limites são definidos para períodos "
    f"mensais, trimestrais e anuais. Quando esses indicadores individuais de continuidade são transgredidos, ou seja, excedem "
    f"o limite estabelecido, a Distribuidora deve compensar financeiramente o consumidor.",

    f"No período analisado, nossos fiscais constataram que houve uma redução dos valores anuais das compensações pagas aos "
    f"consumidores em {ano_final_analise} quando comparado aos anos de 2023 e 2024, isto é, no {localidade2}, houve menos "
    f"interrupções cujos valores ultrapassaram os limites regulatórios definidos pela ANEEL.",

    f"Todavia solicitamos especial apoio à {concessionaria2}, que nos informou por meio da {num_oficio_resp_conc}, em anexo, "
    f"que realizará a partir de 15/12/2025 inspeções nas redes de distribuição de energia elétrica da região. Nessas inspeções "
    f"serão contemplados os trechos dos ramais do {localidade}, de forma a inserir no cronograma de manutenções corretivas e "
    f"preventivas a ser programado para 2026.",

    f"Nesta oportunidade, reiteramos nossos protestos de elevada estima e consideração.",

    f"Atenciosamente,\n\n"
]

palavras_destaque = [f"{pronom_tratamento}", "Ribas do Rio Pardo", f"{localidade}", "Água Clara", "15/12/2025", "2026"]

for texto in paragrafo:
    p = doc.add_paragraph(texto)
    p.style = "Normal"
    destacar_palvras(p, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

    paragrafos = [
        f"CARLOS ALBERTO DE ASSIS",
        f"Diretor-Presidente da Agência Estadual de Regulação de Serviços Públicos de Mato Grosso do Sul\n\n\n\n\n",
        f"Ao {pronom_tratamento}",
        f"{nome_resp_solicitante}",
        f"{cargo_resp_solicitante}",
        f"{comarca}"
    ]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = "Normal"

# === 9. SALVAR DOCX E CONVERTER === #
doc.save(SAIDA_DOCX)
print(f"✅ Relatório DOCX gerado em: {SAIDA_DOCX}")

try:
    convert(SAIDA_DOCX, SAIDA_PDF)
    print(f"✅ Relatório PDF gerado em: {SAIDA_PDF}")
except Exception as e:
    print("⚠️ Erro ao converter para PDF:", e)
