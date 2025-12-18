from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import matplotlib.pyplot as plt
from datetime import date
import os
from docx2pdf import convert
from lxml.html.defs import list_tags
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

# === CONFIGURAÇÕES === #
MUNICIPIO = "Brasilândia"
DATA_RELATORIO = date.today().strftime("%d/%m/%Y")
ANO_RELATORIO = date.today().strftime("%Y")
NUM_RELATORIO = "0XX"

# Caminhos
MODELO = "modelo_agems_demanda.docx"
SAIDA_DOCX = f"saida/relatorio_{MUNICIPIO.lower()}.docx"
SAIDA_PDF = f"saida/relatorio_{MUNICIPIO.lower()}.pdf"

os.makedirs("saida", exist_ok=True)
os.makedirs("figuras", exist_ok=True)

# === ABRE O MODELO BASE === #
doc = Document(MODELO)

# === 1. TÍTULO === #
titulo = doc.add_paragraph(f"Relatório de Análise n° {NUM_RELATORIO}/{ANO_RELATORIO}-CATENE/AGEMS")
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.runs[0]
run.font.size = Pt(14)
run.bold = True

# for style in doc.styles:
    # print(style.name)

# === 2. ASSUNTO === #
solicitante = f"Promotoria de Justiça da Comarca de {MUNICIPIO}, Ministério Público de Mato Grosso do Sul"
num_oficio = "0629/2025/PJ/BRS"
data_oficio = "04 de novembro de 2025"
noticia_fato = "09.2025.00008077-8"
concessionaria = "Energisa Mato Grosso do Sul – EMS"
concessionaria2 = "EMS"
concessionaria3 = "ENERGISA MS"
conjunto1 = "Água Clara"
conjunto2 = conjunto1
ano_inicio_analise = "2022"
mes_inicio_analise = "janeiro"
ano_final_analise = "2025"
mes_final_analise = "novembro"
nivel_tensao_13_8 = "13,8 kV"
nivel_tensao_34_5 = "34,5 kV"
figura = 1
tabela = 1

assunto = (
        f""" \n\n\nAssunto: Atendimento à Demanda da {solicitante}, relativa ao Ofício n° {num_oficio}, que faz referência 
à Notícia de Fato n° {noticia_fato}.\n\n\n
        """
)

a = doc.add_paragraph(assunto)
a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
a.paragraph_format.left_indent = Pt(200)
run1 = a.runs[0]

# === 3. SUMÁRIO AUTOMÁTICO === #
def add_table_of_contents():
    p = doc.add_paragraph("SUMÁRIO\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p.runs[0]
    run2.font.size = Pt(14)
    run2.bold = True
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    run2 = OxmlElement('w:r')
    fldSimple.append(run2)
    p._p.append(fldSimple)

add_table_of_contents()
doc.add_page_break()

def destacar_palvras(paragrafo_obj, texto, palavras_destaque, highlight=True, underline=True):

    texto = paragrafo_obj.text

    paragrafo_obj.clear()  # limpa o parágrafo

    pattern = re.compile("(" + "|".join(re.escape(w) for w in palavras_destaque) + ")")
    partes = re.split(pattern, texto)

    for parte in partes:
        if not parte:
            continue
        run = paragrafo_obj.add_run(parte)
        if parte in palavras_destaque:
            if highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if underline:
                run.underline = True

# === 4. DO OBJETIVO === #
titulo1 = doc.add_paragraph("DO OBJETIVO")
titulo1.style = "Title"

objetivo = doc.add_paragraph(
        f" Subsidiar a Diretoria de Regulação e Fiscalização – Gás Canalizado, Energia e Mineração da AGEMS na resposta" 
        f" ao Ofício n° {num_oficio}, de {data_oficio}, da {solicitante}, referente à solicitação para que a Agência verifique"
        f" possíveis irregularidades e deficiências no fornecimento de energia elétrica por parte da empresa {concessionaria}," 
        f" no município de {MUNICIPIO}."
)
objetivo.style = "No Spacing"

# === 5. DOS FATOS === #
titulo2 = doc.add_paragraph("DOS FATOS")
titulo2.style = "Title"

paragrafos = [
        f" Solicitação da {solicitante} para análise acerca de frequentes quedas de energia elétrica no município de {MUNICIPIO}," 
        f" visando instruir a Notícia de Fato mencionada, em consonância com o Ofício n° 165/2025 da Câmara Municipal de Tacuru'." 
        f" A análise foi elaborada pela Câmara Técnica de Energia e Mineração da AGEMS a partir de informações da ANEEL e da {concessionaria2}."
        f" Nesse contexto, a presente análise foi elaborada pela Câmara Técnica de Energia e Mineração da AGEMS a partir de informações" 
        f" disponibilizadas pela ANEEL e de dados fornecidos pela EMS em resposta ao Ofício da Agência de n° 4732/2025/DGE,"
        f" de 16 de outubro de 2025.",

        f" Neste relatório são detalhados os esclarecimentos da EMS, conforme correspondência ENERGISAMS/DTEC-ANEEL/N°058/2025," 
        f" de 21 de outubro de 2025, com análise das seguintes informações:"
]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = "No Spacing"
    doc.add_paragraph("")

paragrafos = [
    f"  a.	subestações e alimentadores de distribuição que fornecem energia ao município de {MUNICIPIO}, com respectivo "
    f"diagrama unifilar do sistema elétrico e conjuntos aos quais pertencem os consumidores;",
    f"  b.	ocorrências com interrupção na rede de energia elétrica que afetaram os consumidores do município, no período "
    f"de 01/01/2022 a 31/08/2025;",
    f"  c.	valores das compensações pagas aos consumidores pela extrapolação dos limites de continuidade individuais "
    f"(DIC, FIC, DMIC e DICRI), relativos ao período de janeiro/2022 a agosto/2025;",
    f"  d.	manutenções realizadas em 2024 e até agosto de 2025 nas redes de distribuição e equipamentos que atendem o "
    f"município, e as previsões para o restante do ano e para o ano de 2026;",
    f"  e.	obras de melhoria realizadas em 2024 e até agosto de 2025 no município, que tiveram por objetivo a melhora "
    f"dos indicadores de continuidade das unidades consumidoras, como também as previstas para o restante do ano e para "
    f"o ano de 2026.\n"

]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'Heading 7'
    paragrafos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# === 6. DA ANÁLISE === #
titulo3 = doc.add_paragraph("DA ANÁLISE")
titulo3.style = "Title"

#--- III.1 CONTINUIDADE - DEC E FEC --- #

titulo3_1 = doc.add_paragraph(f"Indicadores Coletivos de Continuidade (DEC e FEC) do Município de {MUNICIPIO}")
titulo3_1.style = "Heading 8"

paragrafos = [
    f"O indicador DEC, Duração Equivalente de Interrupção por Unidade Consumidora, é o intervalo de tempo que, em média, no período "
    f"de apuração, em cada unidade consumidora do conjunto considerado, ocorreu descontinuidade da distribuição de energia elétrica.",

    f"O indicador FEC, Frequência Equivalente de Interrupção por Unidade Consumidora, compreende o número de interrupções ocorridas, "
    f"em média, no período de apuração, em cada unidade consumidora do conjunto considerado.",

    f"Os valores dos limites anuais dos indicadores de continuidade dos conjuntos de unidades consumidoras serão disponibilizados "
    f"por meio de audiência pública e estabelecidos em resolução específica, de acordo com a periodicidade da revisão tarifária "
    f"da distribuidora.",

    f"Os valores estabelecidos para o período até a próxima revisão tarifária serão publicados por meio de resolução específica "
    f"e entrarão em vigor a partir do mês de janeiro do ano subsequente à publicação, devendo propiciar melhoria do limite "
    f"anual global de DEC e FEC da distribuidora.",
    
    f"Os gráficos apresentados na Figura {figura}, a seguir, visualizam as evoluções dos Indicadores Anuais de Continuidade "
    f"DEC e FEC dos conjuntos {conjunto1} e {conjunto2} que atendem mais diretamente o município de {MUNICIPIO}, bem como os "
    f"limites permitidos pela ANEEL, no período de {mes_inicio_analise} de {ano_inicio_analise} até {mes_final_analise} de "
    f"{ano_final_analise} (anualizado)."
]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")

# --- GERA E INSERE GRÁFICOS DEC E FEC--- #
df = pd.DataFrame({
    "Causa": ["Vento", "Descarga Atmosférica", "Falha de Equipamento", "Vegetação"],
    "Quantidade": [1290, 1134, 734, 300]
})
plt.figure()
plt.bar(df["Causa"], df["Quantidade"])
plt.title("Causas de Interrupções - 2022 a 2025")
plt.xticks(rotation=15)
plt.tight_layout()
figura1_path = "figuras/figura1.png"
plt.savefig(figura1_path)
plt.close()

doc.add_picture(figura1_path, width=Inches(5.5))
legenda_figura = doc.add_paragraph(f"Figura {figura} - Evolução dos indicadores de continuidade DEC e FEC dos conjuntos {conjunto1} e {conjunto2},"
                                   f"que atendem o município de {MUNICIPIO}, no período {ano_inicio_analise} a {ano_final_analise}\n")
legenda_figura.style = 'Heading 9'
legenda_figura.alignment = WD_ALIGN_PARAGRAPH.CENTER

paragrafos = ["Os gráficos apresentados na Figura 1 mostram que:"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")


paragrafos = [
f"a. Os limites para os indicadores DEC e FEC dos Conjuntos (em cor vermelha) se apresentam em curva descendente, "
f"resultante da exigência da ANEEL por melhora contínua da qualidade de serviço. Dessa forma a cada ano que passa a "
f"distribuidora necessita envidar mais esforços para manter-se dentro dos limites regulados;",

f" b. Dos 2 (dois) conjuntos analisados, as transgressões do indicador anual DEC (em cor azul) do conjunto {conjunto1}" 
f" foram observadas em todos os anos apurados, mostrando um aumento no ano de 2023 e uma pequena queda no ano de 2024. "
f"Já o conjunto {conjunto2} mostra ultrapassagem do limite regulado no ano de 2021, melhorando nos anos seguintes "
f"para valores apurados dentro das metas. Vale destacar que o conjunto de {conjunto1} apresentou maiores números de "
f"ultrapassagens do limite mostrando a necessidade da {concessionaria} investir na região. Para o Conjunto {conjunto2} "
f"observa-se no gráfico que há uma busca da distribuidora em se adaptar aos novos limites estabelecidos pela ANEEL;",

f" c. A respeito das transgressões dos indicadores anuais de FEC (em cor verde), no geral há uma busca de distribuidora "
f" em se adaptar aos novos limites. Somente o conjunto de {conjunto1} apresentou transgressão do valor limite nos "
f"anos de 2020 e 2021, todavia no restante dos anos apurados o indicador se manteve dentro do valor pactuado;",

f" d. Os expurgos por situação de emergência dos indicadores são representados em cor amarela. Os gráficos mostram cada "
f"mês a média móvel dos últimos 12 (doze) meses, dessa forma os eventos climáticos adversos refletem nos expurgos dos "
f"meses seguintes. A tempestade de areia ocorrida em outubro de 2021 impactou os indicadores como pode se perceber nos gráficos;"
]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'Heading 7'
    paragrafos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph("")

#--- III.2 COMPENSAÇÃO - DIC, FIC, DMIC E DICRI --- #

titulo3_2 = doc.add_paragraph(f"Indicadores Individuais de Continuidade - Compensação")
titulo3_2.style = "Heading 8"

paragrafos = [
    f" Os indicadores individuais de continuidade DIC (Duração de Interrupção Individual por Unidade Consumidora), "
    f"FIC (Frequência de Interrupção Individual por Unidade Consumidora) e DMIC (Duração Máxima de Interrupção Contínua "
    f"por Unidade Consumidora) são destacados na fatura de energia elétrica do consumidor. Esses indicadores representam "
    f"para o consumidor a qualidade dos serviços prestados pela distribuidora e mensuram a frequência e a duração das "
    f"interrupções ocorridas em sua unidade.",

    f" Os limites são definidos no Módulo 8 do PRODIST, para períodos mensais, trimestrais e anuais, com base na variação "
    f"dos limites anuais dos indicadores de continuidade coletivos (DEC e FEC) dos conjuntos elétricos de consumidores.",

    f" Quando esses indicadores individuais de continuidade são transgredidos, ou seja, excedem o limite estabelecido, "
    f"a distribuidora deve compensar financeiramente o consumidor. A compensação é automática, e deve ser paga em até 2 "
    f"(dois) meses após o mês de apuração do indicador (mês em que houve a interrupção).",

    f" A Figura {figura}, a seguir apresenta os montantes de compensações pagas por violação aos limites dos indicadores "
    f"individuais aos consumidores do município de {MUNICIPIO} no período de {ano_inicio_analise} até {mes_final_analise} "
    f"de {ano_final_analise}."
]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")

#INSERIR Gráfico de COMPENSAÇÃO#

legenda_figura = doc.add_paragraph(f"Figura {figura} - Montantes de compensações anuais creditados aos consumidores por "
                                   f"violação dos limites dos indicadores individuais no município {MUNICIPIO},no período "
                                   f"de {ano_inicio_analise} até {ano_final_analise}\n")
legenda_figura.style = 'Heading 9'
legenda_figura.alignment = WD_ALIGN_PARAGRAPH.CENTER
figura = figura + 1

paragrafos = [
    f"Percebe-se que houve um aumento do valor pago nos anos de 2022 e 2023 em relação a 2021. Em 2024, os dados das "
    f"compensações pagas são referentes ao período de janeiro a maio. Ao comparar com o mesmo período dos anos anteriores "
    f"de 2022 e 2023, o valor de 2024 demonstra uma tendência de alcançar os mesmos patamares anteriores, indicando que "
    f"os problemas perduram no município.",

    f"O aumento das compensações aos consumidores é resultante de limites mais rigorosos exigidos para a qualidade da "
    f"distribuição de energia elétrica. A compensação diretamente na fatura do consumidor por transgressão de limites de "
    f"indicadores é peça regulatória efetiva. É uma obrigação imposta à distribuidora que independe de uma ação "
    f"fiscalizatória predeterminada. É um sinal regulatório cujo objetivo é fazer com que a concessionária invista na "
    f"qualidade do serviço."
]

palavras_destaque = ["aumento", "2022", "2021", "2023", "2024"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

#--- III.3 MANUTENÇÕES E OBRAS PREVISTAS --- #

titulo3_3 = doc.add_paragraph(f"Manutenções e Obras Previstas")
titulo3_3.style = "Heading 8"

paragrafos = [
    f"Foi requisitado à concessionária {concessionaria2} informar as manutenções realizadas em {ano_final_analise} no "
    f"município e as previsões para o ano de 2024, conforme mostrado na Tabela {tabela} a seguir."
]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")

#INSERIR TABELA MANUTENÇÕES#

legenda_tabela = doc.add_paragraph(f"Tabela {tabela} - Dados de valores executados e previstos de {ano_final_analise} dos serviços "
                                   f"de poda de árvore, manutenção e limpeza por faixa informados pela concessionária, "
                                   f"referente ao município de {MUNICIPIO}\n")
legenda_tabela.style = 'Heading 9'
legenda_tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
tabela = tabela + 1

paragrafos = [
     f"Pela tabela fica evidente que há intervenção da concessionária a respeito de podas de árvores, buscando trazer "
     f"maior eficácia ao sistema de distribuição, sendo já executado 80,2% do previsto para o ano de 2024.",

     f"Com relação as manutenções por estrutura e as limpezas de faixa de servidão nas linhas e redes de distribuição, "
     f"a {concessionaria3} tem programações para {ano_final_analise}, com maior previsão para limpeza de faixa que é mais "
     f"atinente a linha de distribuição de {nivel_tensao_34_5}.",
    
     f"Além das manutenções informadas, a {concessionaria} comunicou a realização de obras de substituição de cabos, postes, chaves, "
     f"transformadores e equipamentos de rede de distribuição nos alimentadores que atendem a região. Porém, atualmente "
     f"não está previsto nem uma obra estruturante para o ano de {ano_final_analise}.",
    
     f"A equipe de fiscalização da AGEMS fez uma inspeção in loco no município de {MUNICIPIO}, no dia 26/06/2024, "
     f"onde constatou a realização das manutenções e melhorias e o estado de conservação das redes de distribuição "
     f"primária e secundária, não verificando irregularidades significativas. No Anexo deste relatório constam fotos da "
     f"visita apresentando pequenas irregularidades, relativo somente a necessidade da poda de galhos de árvores próximos "
     f"as redes, que inclusive devem estar na previsão para o segundo semestre de 2024."
]

palavras_destaque = ["AGEMS", "2024", "26/06/" ]

for texto in paragrafos:
    p = doc.add_paragraph(texto)
    p.style = 'No Spacing'
    destacar_palvras(p, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

#--- III.4 Interrupções na Rede de Energia Elétrica que Afetaram os Consumidores da Região --- #

titulo3_4 = doc.add_paragraph(f"Interrupções na Rede de Energia Elétrica que Afetaram os Consumidores da Região")
titulo3_4.style = "Heading 8"

paragrafos = [
     f"Pelos dados das interrupções de longa duração apresentados pela {concessionaria3}, referentes ao período de {ano_inicio_analise} a "
     f"{ano_final_analise}, a Tabela {tabela} mostra as causas que impactaram nas ocorrências não programadas."

]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")

df = pd.DataFrame({
    "Causa": ["Vento", "Descarga Atmosférica", "Falha de Equipamento", "Vegetação"],
    "Quantidade": [1290, 1134, 734, 300]
})
plt.figure()
plt.bar(df["Causa"], df["Quantidade"])
plt.title("Causas de Interrupções - 2022 a 2025")
plt.xticks(rotation=15)
plt.tight_layout()
figura1_path = "figuras/figura1.png"
plt.savefig(figura1_path)
plt.close()

doc.add_picture(figura1_path, width=Inches(5.5))

legenda_tabela = doc.add_paragraph(f"Tabela {tabela} - Dados das interrupções não programadas do período de {ano_inicio_analise} a {ano_final_analise}\n")
legenda_tabela.style = 'Heading 9'
legenda_tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
tabela = tabela + 1

paragrafos = [
    f"Como pode ser observado na tabela, dos 5.581 registros de interrupções causadas por desligamentos não programados, as "
    f"causas que impactaram nas ocorrências foram: 3.155 por Descarga Atmosférica, 1.663 por Ventos, 762 devido Árvore ou "
    f"Vegetação e apenas uma ocorrência por Erosão."
]

palavras_destaque = ["5.581","causas que impactaram nas ocorrências foram"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

#--- III.5 Interrupções na Rede de Energia Elétrica que Afetaram os Consumidores da Região --- #

titulo3_5 = doc.add_paragraph(f"Tempo Médio de Atendimento às Ocorrências Emergenciais - TMAE")
titulo3_5.style = "Heading 8"

paragrafos = [
    f"Foi constatado, no período de 01/01/2021 a 30/04/2024, que os tempos de atendimento, desde a reclamação ou "
    f"constatação da EMS da ocorrência até a restauração do fornecimento de energia elétrica, são motivo de preocupação, "
    f"visto que, o indicador TMAE desse período é de 12,66 horas.",

    f"Com relação aos tempos de atendimento total, desde o conhecimento/reclamação até a restauração, foi apurado que no "
    f"ano de 2021 o tempo médio de atendimento a emergências foi de 16,33 horas, em 2022 foi de 10,51 horas, em 2023 foi "
    f"de 11,66 horas e no período de janeiro a abril de 2024 foi de 10,95 horas, conforme a Figura 3."

]

palavras_destaque = ["01/01/2021 a 30/04/2024", "12,66"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

legenda_figura = doc.add_paragraph(f"Figura {figura} - Valores anuais do TMAE (Tempo Médio de Atendimento a Emergências) "
                                   f"em horas no período de 2021, 2022, 2023 e de janeiro a abril de 2024 \n")
legenda_figura.style = 'Heading 9'
legenda_figura.alignment = WD_ALIGN_PARAGRAPH.CENTER
figura = figura + 1

paragrafos = [
    f"Considerando os critérios de fiscalização, tempos médios anuais de interrupção de energia superiores a 6 horas já "
    f"são preocupantes. Em situação de alerta estão os tempos entre 6 e 12 horas, enquanto aqueles acima de 12 horas indicam "
    f"irregularidade. Constata-se que os valores para a região se encontram em “situação de alarme.",

    f"Por mais que houve uma queda do indicador referente a 2021, os outros períodos demonstram valores elevados, e ao "
    f"analisar o ano de 2024, percebe-se que a tendência é alcançar o mesmo patamar dos valores anteriores.",

    f"Neste período de 2023 a maio de 2024 ocorreram 336 interrupções com tempo de restabelecimento superior a 24 horas.",

    f"Na Figura 4 a seguir é demonstrado o tempo de atendimento e a quantidade de eventos não programados do total das interrupções."

]

palavras_destaque = ["Neste período de 2023 a maio de 2024 ocorreram 336 interrupções com tempo de restabelecimento superior a 24 horas"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

legenda_figura = doc.add_paragraph(f"Figura {figura} - Número de interrupções relacionadas com a quantidade de horas sem "
                                   f"energia no período de 2023 até abril de 2024 \n")
legenda_figura.style = 'Heading 9'
legenda_figura.alignment = WD_ALIGN_PARAGRAPH.CENTER
figura = figura + 1

#--- III.6 Interrupções na Rede de Energia Elétrica que Afetaram os Consumidores da Região --- #

titulo3_6 = doc.add_paragraph(f"Pedidos de Ressarcimento de Danos Elétricos de Consumidores")
titulo3_6.style = "Heading 8"

paragrafos = [
    f"Os pedidos de ressarcimento são referentes a queima de eletrodomésticos e equipamentos eletroeletrônicos.",

    f"Em atendimento a solicitação da AGEMS, a Energisa MS relacionou os pedidos de ressarcimento de danos elétricos dos "
    f"consumidores do município, no período de 01/01/2021 a 31/05/2024.",

    f"Foi informado um total de 45 pedidos sendo 4 procedentes e 41 improcedentes. Para os improcedentes, foi informado "
    f"o motivo do indeferimento.",

    f"Portanto, os pedidos procedentes corresponderam somente a 8,88% do total.",

    f"Além disso, por conta da situação da distribuição de energia elétrica encontrada no {MUNICIPIO}, o promotor anexou ao "
    f"ofício em questão, um abaixo assinado contendo 533 assinaturas com comentários a respeito da qualidade do serviço "
    f"prestado pela concessionária.",

    f"A maior parte das reclamações estão voltadas para os seguintes temas:"

]

palavras_destaque = ["45 pedidos sendo 4 procedentes e 41 improcedentes", "553"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

paragrafos = [
    f" a. Continuidade;",
    f" b. Danos/perda de eletrodomésticos e eletroeletrônicos;",
    f" c. Queda de energia proveniente de intempéries climáticas;",
    f" d. Demora na religação de energia;",
    f" e. Atendimento do Call Center ineficiente;"

]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafos.style = 'Heading 7'

doc.add_paragraph("")

#--- III.7 Problemas relacionados ao nível de tensão (Precário e Crítico) --- #
titulo3_7 = doc.add_paragraph(f"Problemas Relacionados ao Nível de Tensão (Precário e Crítico)")
titulo3_7.style = "Heading 8"

paragrafos = [
    f"Quanto aos níveis de tensão, que representam a qualidade do produto entregue aos consumidores, a ANEEL determina "
    f"que as concessionárias de energia elétrica mantenham um padrão de conformidade da tensão em regime permanente, "
    f"estabelecendo os limites adequados, precários e críticos, os indicadores de controle, os critérios de medição e de "
    f"registro e os prazos para compensação ao consumidor, caso as medições de tensão excedam os limites dos indicadores.",

    f"Os indicadores DRP (duração relativa da transgressão de tensão precária) e DRC (duração relativa da transgressão "
    f"de tensão crítica) expressam o percentual do tempo em que determinada unidade consumidora permanece com tensão "
    f"precária e com tensão crítica, conforme definido no Módulo 8 dos Procedimentos de Distribuição de Energia Elétrica "
    f"no Sistema Elétrico Nacional – PRODIST.",

    f"Além dos pedidos de verificação dos níveis de tensão solicitados pelos consumidores (reclamações), a distribuidora "
    f"deve efetuar também medições amostrais de tensão conforme critérios estabelecidos no Módulo 8 do PRODIST. O sorteio "
    f"da amostra das unidades consumidoras de cada distribuidora para fins de medição será realizado pela ANEEL, no mês "
    f"de outubro de cada ano, por meio de critério estatístico aleatório, a partir das Bases de Dados Geográficas das "
    f"Distribuidoras - BDGD.",

    f"Caso as medições de tensão indiquem valor de DRP superior ao DRPlimite, ou valor de DRC superior ao DRClimite, "
    f"estabelecidos no PRODIST (DRPLimite: 3%, e DRCLimite: 0,5%.), a distribuidora deve regularizar a tensão de atendimento, "
    f"sem prejuízo do pagamento de compensação aos consumidores e das sanções cabíveis pela fiscalização da ANEEL.",

    f"A compensação deve ser mantida enquanto o(s) indicador(es) DRP e/ou DRC for(em) superior(es) aos limites estabelecidos. "
    f"O valor da compensação deve ser creditado na fatura emitida no prazo máximo de 2 meses subsequentes ao mês civil de "
    f"referência da última medição que constatou a violação.",

    f"Foi solicitado à {concessionaria} os pedidos de verificação de tensão dos consumidores do {MUNICIPIO} no período de"
    f"janeiro de 2022 a junho de 2025, a relação das eventuais medições amostrais no referido bairro (se ocorreram) e as"
    f"eventuais amostrais nas unidades consumidoras localizadas no Conjunto Elétrico {conjunto1}.",
    
    f"A {concessionaria2} informou que ocorreram 16 pedidos de medição de tensão referente a reclamações de consumidores, "
    f"no período de janeiro/2022 a junho/2025, sendo que em 15 os níveis de tensão estavam dentro dos limites adequados. "
    f"Em uma das solicitações foi constatada violação de tensão, com a concessionária tomando as providências necessárias "
    f"e o processo de regularização concluído e comprovado com a realização de medição de 7 dias.",

    f"Quanto as medições amostrais, a distribuidora informou um caso de um consumidor do bairro, com a medição efetuada e "
    f"constatado nível de tensão adequado. Para as amostrais de demais consumidores do Conjunto {conjunto1} foram 5 casos, "
    f"constatando níveis de tensão dentro dos padrões adequados em 4 casos e um caso fora dos limites adequados que está sendo "
    f"regularizado pela EMS nos próximos dias."

]

palavras_destaque = ["janeiro de 2022 a junho de 2025","Conjunto Elétrico","16"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

#--- III.8 Nível de carregamento dos transformadores --- #
titulo3_8 = doc.add_paragraph(f"Nível de Carregamento dos Transformadores")
titulo3_8.style = "Heading 8"

paragrafos = [
    f"A {concessionaria3} encaminhou a relação dos transformadores de distribuição que atendem as unidades consumidoras da região com"
    f"as devidas potências em kVA e o carregamento percentual que permite verificar se há sobrecarga nos mesmos e eventual"
    f"necessidade de substituição. Os valores apresentados dos carregamentos foram calculados por modelo matemático que"
    f"considera principalmente o pico de consumo dos consumidores nos últimos 12 meses.",
    
    f"Este modelo também considera uma margem aceitável de carregamento, sendo que o único transformador que está com "
    f"carregamento acima de 100%, TD 59295 de 10 kVA que apresenta 110,28% no alimentador CGS08, está dentro dessa margem "
    f"e não há registros de problemas no circuito (aberturas) por sobrecarga.",
    
    f"Feito a inspeção do transformador TD 59295 “in loco” pelos fiscais da AGEMS, verificou não haver carga ou previsão "
    f"de aumento que justifique um maior carregamento do mesmo. Não há registros de problemas no circuito (aberturas) por "
    f"sobrecarga."

]

palavras_destaque = ["único","TD 59295"]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    destacar_palvras(paragrafos, texto, palavras_destaque, highlight=True, underline=True)
    doc.add_paragraph("")

# === 7. CONCLUSÃO === #

titulo4 = doc.add_paragraph("DA CONCLUSÃO")
titulo4.style = "Title"

paragrafos = [
    f"As questões encaminhadas pela Promotoria de Justiça da Comarca de Nova Alvorada do Sul/MS foram analisadas com base "
    f"nas informações estabelecidas nos regulamentos e disponibilizadas no sítio Aneel, e nas fiscalizações realizadas "
    f"pela AGEMS no Âmbito dos contratos de Metas SFE/AGEMS.",

    f"Para o município Nova Alvorado do Sul foram apresentados os gráficos dos conjuntos elétricos principais que compõe "
    f"a região geoelétrica do município. Os limites para os indicadores se apresentam em curva descendente, resultante da "
    f"exigência da ANEEL por melhora contínua da qualidade de Serviço. Dessa forma a cada ano que passa a distribuidora "
    f"necessita envidar mais esforços para manter-se dentro dos limites regulado.",

    f"Referente ao indicador DEC, relativo a duração das interrupções, dos conjuntos analisados, para o conjunto Rio "
    f"Brilhante as transgressões do limite foram observadas em todos os períodos de análise, porém, o gráfico apresenta "
    f"uma queda dos valores anualizados apurados nos primeiros meses de 2024 próximo do limite estabelecido pela ANEEL. "
    f"Para o conjunto Rio Brilhante Rural, houve transgressão somente no ano de 2021, mas desde então o indicador se "
    f"estabeleceu dentro do limite.",

    f"Ressalta-se a necessidade de atenção da ENERGISA MS com relação ao conjunto Rio Brilhante, procurando manter "
    f"programações constantes de manutenções por estrutura, podas de árvores e limpeza de faixa de servidão. Deve-se "
    f"dar prioridade às manutenções preventivas nas duas linhas de distribuição em 34,5 kV de Rio Brilhante a "
    f"Nova Alvorada do Sul.",

    f"Para o indicador FEC, relativo a quantidade de interrupções, os gráficos sinalizam valores apurados dentro das metas.",
    
    f"Com relação aos indicadores individuais, DIC, FIC e DMIC, a EMS vem realizando as compensações pelas ultrapassagens "
    f"dos limites regulados. Os valores apurados dos limites individuais e as respectivas compensações tendem a diminuir "
    f"com a melhora dos indicadores de continuidade coletivos (DEC e FEC).",
    
    f"A AGEMS, além das fiscalizações programadas nos contratos de metas com a ANEEL, realiza eventos em parceria com a "
    f"Distribuidora Energisa para um público-alvo específico, obtendo apoio institucional para o plantio adequado, manejo "
    f"e poda de árvores, atividades essenciais na redução de interrupções do fornecimento de energia e prevenção de "
    f"acidentes, tendo realizado um evento em Dourados no dia 16 de maio de 2024 e estando programado para o dia 01 de "
    f"agosto de 2024, em Amambai, destinado a representantes dos municípios do cone sul de Mato Grosso do Sul.",
    
    f"Por fim, no âmbito do Convênio de Cooperação com a ANEEL, a AGEMS iniciou em fevereiro de 2024 uma campanha de "
    f"fiscalizações em subestações e alimentadores da Energisa MS que se estenderá ao longo deste ano de 2024 e serão "
    f"contemplados os alimentadores com mais ocorrências de interrupção de energia.\n\n\n\n"

]

for texto in paragrafos:
    paragrafos = doc.add_paragraph(texto)
    paragrafos.style = 'No Spacing'
    doc.add_paragraph("")

cg = doc.add_paragraph(f"Campo Grande, {DATA_RELATORIO}\n\n\n\n")
cg.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = cg.runs[0]
run.font.size = Pt(12)


# === 9. SALVAR DOCX E CONVERTER === #
doc.save(SAIDA_DOCX)
print(f"✅ Relatório DOCX gerado em: {SAIDA_DOCX}")

try:
    convert(SAIDA_DOCX, SAIDA_PDF)
    print(f"✅ Relatório PDF gerado em: {SAIDA_PDF}")
except Exception as e:
    print("⚠️ Erro ao converter para PDF:", e)
